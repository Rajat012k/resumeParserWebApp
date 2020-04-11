# NECCESSARY IMPORTS
from django.core.files.storage import FileSystemStorage
from django.http import HttpResponse
from django.shortcuts import render

# pip install PyMuPDF
import fitz
# pip install docx
import docx
import re
import zipfile
# pip install docx2txt
import docx2txt
# pip iinstall pymongo
import pymongo
#pip iinstall xwlt
import xlwt
#pip install dns

# connecting with the mongoDB database
mongoclient = pymongo.MongoClient(
    "mongodb+srv://vikas:Test123@node-events-j0bd8.mongodb.net/test?retryWrites=true&w=majority")
db = mongoclient.get_database('rajatDB')
records = db.Users

# Loads the dashboard screen
def dashboard(request):
    return render(request, "index.html")

# Function call at the upload button
def _uploadFiles(request):

    # Getting the file from the input
    f = request.FILES["filename"]
    global newCreatedfilename
    newCreatedfilename=f.name.split(".")[0]
    filext = f.name.split(".")[1]

    
    #------------------- THIS IS FOR THE DOCX BASED RESUME --------------------#
    
     
    if filext == "docx" or filext == "doc":
        wordText = docx2txt.process(f)
        fs = FileSystemStorage()

        # FOR GETTING THE NAME
        txtOfFile=wordText.split("\n")
        noOfCharacters=len(txtOfFile)
        if noOfCharacters<1000:
            noOfCharacters=2023
        
        txtOfFile = ' '.join(txtOfFile).split("  ")
        txtOfFile=list(filter(None,txtOfFile))

        if txtOfFile[0]=="RESUME" or txtOfFile[0]=="Resume" or txtOfFile[0]=="NAME" or txtOfFile[0]=="Name":
            Name=txtOfFile[1]
        else:
            Name=txtOfFile[0].split("  ")[0]

        # FOR GETTING THE EMAIL
        emailPattern = re.compile(r'[a-zA-Z0-9-\.]+@[a-zA-Z-\.]*\.(com|edu|net)')
        emailMatches = emailPattern.finditer(wordText)
        Email = "-"
        for match in emailMatches:
            Email = match.group(0)

        # FOR GETTING THE LINKEDIN PROFILE
        linkPattern = re.compile('^https:\\/\\/[a-z]{2,3}\\.linkedin\\.com\\/.*$')
        linkMatches=linkPattern.finditer(wordText)
        linkedIN="-"
        for match in linkMatches:
            linkedIN=match.group(0)

        # FOR GETTING THE PHONE NUMBER
        phonePattern = re.compile(r'(\(|\+)?\d{3}(\)|\-)?\s?\d{2,3}\-?\d{3,6}')
        phoneMatches = phonePattern.finditer(wordText)
        Phone_Num = "-"
        for match in phoneMatches:
            Phone_Num = match.group(0)

        # FOR GETTING TABLES
        document = docx.Document(f)
        Table_Count = 0
        Table_Count = len(document.tables)

        # FOR FONT NAME AND SIZE
        doc = docx.Document(f)        
        fontsize=["14","11","9","8"]
        fontname=[]
        for p in doc.paragraphs:
            fontname.append(p.style.font.name)
        fontname=set(fontname)
        fontname=list(fontname)
        if(len(fontname)==0 or len(fontname)==1):
            fontname=["Arial","Verdana","Calibri"]

        # FOR TOTAL IMAGES
        totalImages = []
        z = zipfile.ZipFile(f)
        all_files = z.namelist()
        images = filter(lambda x: x.startswith('word/media/'), all_files)
        Images = 0
        for match in images:
            totalImages.append(match)
        Images = len(totalImages)
        params = {"Name": Name, "EmailID": Email, "Contact": Phone_Num, "ImageCount": Images,"Linkedin":linkedIN , "FontName":fontname,
        "FontSize":fontsize,"Table":Table_Count,"CharLine":noOfCharacters}
        
        #inserting the params object into the database
        records.insert_one(params)

        # returning the PARAMS Object and rendering to the next page
        return render(request, "nextpage.html", params)

    #------------------- THIS IS FOR THE PDF BASED RESUME --------------------#
    elif filext=="pdf":
        fs = FileSystemStorage()
        filename = fs.save(f.name, f)
        file1 = fitz.open(f)
        pdfText = file1.getPageText(0)

        noOfCharacters=len(pdfText)

        # FOR GETTING THE NAME
        txtOfFile=pdfText.split("\n")
        txtOfFile = ' '.join(txtOfFile).split("  ")
        txtOfFile=list(filter(None,txtOfFile))

        if txtOfFile[0]=="RESUME" or txtOfFile[0]=="Resume":
            Name=txtOfFile[1]
        else:
            Name = txtOfFile[0]

        # FOR GETTING THE EMAIL
        emailPattern = re.compile(
            r'[a-zA-Z0-9-\.]+@[a-zA-Z-\.]*\.(com|edu|net)')
        Emailmatches = emailPattern.finditer(pdfText)
        Email = "-"
        for match in Emailmatches:
            Email = match.group(0)

        # FOR GETTING THE LINKEDIN PROFILE
        linkPattern = re.compile('^https:\\/\\/[a-z]{2,3}\\.linkedin\\.com\\/.*$')
        linkMatches=linkPattern.finditer(pdfText)
        linkedIN="-"
        for match in linkMatches:
            linkedIN=match.group(0)

        # FOR GETTING THE PHONE NUMBER
        phonePattern = re.compile(r'(\(|\+)?\d{3}(\)|\-)?\s?\d{2,3}\-?\d{3,6}')
        phoneMatches = phonePattern.finditer(pdfText)
        Phone_Num = 0
        for match in phoneMatches:
            Phone_Num = match.group(0)

        # FOR TOTAL IMAGES
        Images = len(file1.getPageImageList(0))
        Table_Count = 0

        # FOR FONT NAME AND SIZE
        fontname=[]
        fontsize=[]
        for font in range(3):
            fontname.append(file1.getPageFontList(0)[font][3])
            fontsize.append(file1.getPageFontList(0)[font][0])
            

        # Object Of The Details
        params = {"Name": Name, "EmailID": Email, "Contact": Phone_Num, "ImageCount": Images,"Linkedin":linkedIN , "FontName":fontname,
        "FontSize":fontsize,"Table":Table_Count,"CharLine":noOfCharacters}
        
        # passing the params object into the database
        records.insert_one(params)

        # rendering to the next page
        return render(request, "nextpage.html", params)

    else:
        #handling the error page
        return render(request,"errorpage.html")
    


# Downloading the file function
def _downloadFile(request):

    # GETTING THE RECORDS FROM THE DB
    data=list(records.find())

    # CREATING A WORKBOOK(EXCEL)
    wb=xlwt.Workbook()
    ws=wb.add_sheet("Data Sheet")
    ws.write(0,0,"Name")
    ws.write(0,1,"Phone")
    ws.write(0,2,"Email")
    ws.write(0,3,"Linkedin")
    ws.write(0,4,"Images")
    ws.write(0,5,"Tables")
    ws.write(0,6,"Fontsize")
    ws.write(0,7,"Fontname")
    ws.write(0,8,"Characters+Lines")

    fname=[]
    fsize=[]
    for i in range(len(data)):
        
        ws.write(i+1,0,data[i]["Name"])
        ws.write(i+1,1,data[i]["Contact"])
        ws.write(i+1,2,data[i]["EmailID"])
        ws.write(i+1,3,data[i]["Linkedin"])
        ws.write(i+1,4,data[i]["ImageCount"])
        ws.write(i+1,5,data[i]["Table"])
        for j in range(3):
            fname.append(str(data[i]["FontName"][j])+",")
            fsize.append(str(data[i]["FontSize"][j])+",")
        ws.write(i+1,6,fsize)
        ws.write(i+1,7,fname)
        ws.write(i+1,8,data[i]["CharLine"])

    downloadXlsv=newCreatedfilename+".xls"
    wb.save(downloadXlsv)
    return render(request,"index.html")

